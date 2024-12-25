import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os

folder_path = r"C:\Users\Rithesh\Downloads\Intern\Excel"

def save_eda_visualizations(df, folder_path):
    image_filenames = []
    
    hist_file = os.path.join(folder_path, "data_distribution.png")
    plt.figure(figsize=(12, 8))
    columns = ['current_price', 'market_cap', 'total_volume', 'price_change_percentage_24h']
    colors = sns.color_palette("Set2", len(columns))
    
    for i, col in enumerate(columns):
        plt.subplot(2, 2, i + 1)
        df[col].hist(bins=20, color=colors[i], edgecolor='black')
        plt.title(f"Distribution of {col}", fontsize=14)
        plt.xlabel(col, fontsize=12)
        plt.ylabel("Frequency", fontsize=12)
    
    plt.tight_layout()
    plt.savefig(hist_file)
    image_filenames.append(hist_file)
    plt.close()
    
    corr_file = os.path.join(folder_path, "correlation_matrix.png")
    corr_matrix = df[['current_price', 'market_cap', 'total_volume', 'price_change_percentage_24h']].corr()
    plt.figure(figsize=(10, 6))
    sns.heatmap(corr_matrix, annot=True, cmap='RdYlBu', fmt='.2f', linewidths=0.5)
    plt.title("Correlation Matrix", fontsize=16)
    plt.savefig(corr_file)
    image_filenames.append(corr_file)
    plt.close()
    

    bar_file = os.path.join(folder_path, "top_5_by_market_cap.png")
    top_5 = df.nlargest(5, 'market_cap')[['name', 'market_cap']]
    plt.figure(figsize=(10, 6))
    sns.barplot(x='market_cap', y='name', data=top_5, palette='coolwarm')
    plt.title("Top 5 Cryptocurrencies by Market Cap", fontsize=16)
    plt.xlabel('Market Cap (USD)', fontsize=14)
    plt.ylabel('Cryptocurrency', fontsize=14)
    plt.savefig(bar_file)
    image_filenames.append(bar_file)
    plt.close()
    
    scatter_file = os.path.join(folder_path, "price_vs_market_cap.png")
    plt.figure(figsize=(10, 6))
    sns.scatterplot(x='current_price', y='market_cap', data=df, hue='price_change_percentage_24h', palette='magma', s=100, edgecolor='black')
    plt.title("Price vs Market Cap", fontsize=16)
    plt.xlabel('Current Price (USD)', fontsize=14)
    plt.ylabel('Market Cap (USD)', fontsize=14)
    plt.savefig(scatter_file)
    image_filenames.append(scatter_file)
    plt.close()
    
    return image_filenames

def generate_report(data, folder_path):
    top_5 = data.nlargest(5, 'market_cap')[['name', 'market_cap']]
    average_price = data['current_price'].mean()
    highest_change = data.loc[data['price_change_percentage_24h'].idxmax()]
    lowest_change = data.loc[data['price_change_percentage_24h'].idxmin()]

    doc = Document()
    doc.add_heading('Cryptocurrency Analysis Report', level=1)

    doc.add_heading('1. Top 5 Cryptocurrencies by Market Capitalization', level=2)
    for _, row in top_5.iterrows():
        doc.add_paragraph(f"{row['name']}: ${row['market_cap']:,}")

    doc.add_heading('2. Average Price of Top 50 Cryptocurrencies', level=2)
    doc.add_paragraph(f"The average price of the top 50 cryptocurrencies is ${average_price:,.2f}.")

    doc.add_heading('3. 24-hour Price Change Analysis', level=2)
    doc.add_paragraph(
        f"Cryptocurrency with the highest 24-hour percentage change: "
        f"{highest_change['name']} ({highest_change['symbol']}): {highest_change['price_change_percentage_24h']:.2f}%."
    )
    doc.add_paragraph(
        f"Cryptocurrency with the lowest 24-hour percentage change: "
        f"{lowest_change['name']} ({lowest_change['symbol']}): {lowest_change['price_change_percentage_24h']:.2f}%."
    )

    return doc

def add_eda_to_report(doc, df, folder_path):
    image_filenames = save_eda_visualizations(df, folder_path)
    doc.add_heading('4. EDA Visualizations', level=2)
    for image_file in image_filenames:
        doc.add_paragraph(f"Figure: {image_file}")
        doc.add_picture(image_file, width=Inches(6)) 
        doc.add_paragraph("\n")
    
    return doc

def create_and_save_complete_report(data, folder_path):
    doc = generate_report(data, folder_path)
    doc = add_eda_to_report(doc, data, folder_path)
    
    
    report_path = os.path.join(folder_path, "cryptocurrency_analysis_complete_report.docx")
    doc.save(report_path)
    print(f"Complete report saved at: {report_path}")


crypto_data = fetch_crypto_data()  
if crypto_data is not None:
    create_and_save_complete_report(crypto_data, folder_path)
